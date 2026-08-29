"""
Generate a professional business plan (.docx) for EU StudyBridge Nigeria.
Based on the services, routes, and positioning shown in the website (index.html / script.js).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x10, 0x3A, 0xC0)   # brand blue
DARK = RGBColor(0x10, 0x18, 0x28)      # near-black
MUTED = RGBColor(0x52, 0x60, 0x7A)
GREEN = RGBColor(0x11, 0xA5, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = PRIMARY
h1.paragraph_format.space_before = Pt(16)
h1.paragraph_format.space_after = Pt(6)

h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = DARK
h2.paragraph_format.space_before = Pt(10)
h2.paragraph_format.space_after = Pt(4)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def para(text, size=11, bold=False, italic=False, color=DARK, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
    p.add_run(text)
    return p


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = PRIMARY if level == 1 else DARK
        r.font.name = 'Calibri'
    return h


def add_table(headers, rows, widths=None, header_fill='103AC0'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ==================== COVER PAGE ====================
for _ in range(4):
    doc.add_paragraph()
para('BUSINESS PLAN', size=13, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para('EU StudyBridge Nigeria', size=34, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('Study Abroad Consulting — Connecting Nigerian Students to European Universities', size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

cover_items = [
    ('Prepared for', 'Founders & Management Team'),
    ('Prepared by', 'Strategy & Business Development'),
    ('Date', 'August 2026'),
    ('Office', '14 Ikoyi Road, Lekki Phase 1, Lagos, Nigeria'),
    ('Contact', 'hello@eustudybridge.ng  |  +234 803 440 3121'),
    ('Website', 'eustudybridge.ng'),
]
for label, value in cover_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label}:  ')
    r.font.bold = True
    r.font.color.rgb = MUTED
    r.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ==================== TABLE OF CONTENTS (manual) ====================
heading('Contents', level=1)
toc = [
    '1. Executive Summary',
    '2. Company Overview',
    '3. Market Analysis',
    '4. Services & Revenue Model',
    '5. Marketing & Sales Strategy',
    '6. Operations Plan',
    '7. Management & Organization',
    '8. Financial Plan',
    '9. Funding Requirements',
    '10. Risks & Mitigation',
    '11. Implementation Roadmap',
    'Appendix A: Service Fee Schedule',
    'Appendix B: 24-Month Financial Projection',
]
for t in toc:
    para(t, size=12, bold=False, color=DARK, space_after=4)
doc.add_page_break()

# ==================== 1. EXECUTIVE SUMMARY ====================
heading('1. Executive Summary', level=1)
para(
    'EU StudyBridge Nigeria is a study-abroad consultancy based in Lagos, Nigeria, that guides Nigerian '
    'students through every step of studying in the European Union — from university selection and admission '
    'applications to visa strategy, scholarship matching, and pre-departure preparation. The business addresses '
    'a fast-growing demand: tens of thousands of Nigerian students apply to EU universities each year, yet most '
    'lack reliable guidance, structured timelines, and document preparation support.'
)
para(
    'The company operates a low-fixed-cost, high-margin advisory model. Revenue comes from per-student service '
    'fees across four study routes (Undergraduate, Master\u2019s, Language + Foundation, and Doctoral/Research) '
    'and from destination countries including Germany, Ireland, the Netherlands, Poland, Italy, and Spain. '
    'Targeting 240 enrolled students by the end of year two, the business is projected to reach '
    '₦186.4 million in annual revenue with a 49% net margin in Year 2.'
)
para(
    'Initial funding of ₦18 million covers team, office, marketing, and working capital for the first 12 months. '
    'EU StudyBridge is positioned to become the most trusted study-abroad partner for Nigerian students through '
    'personalised, transparent, and outcome-focused service.'
)

# ==================== 2. COMPANY OVERVIEW ====================
heading('2. Company Overview', level=1)
heading('2.1 Mission', level=2)
para(
    'To make European higher education accessible to Nigerian students through expert, personalised, and '
    'transparent guidance at every stage of their journey.'
)
heading('2.2 Vision', level=2)
para(
    'To be the most trusted bridge between Nigeria and European universities — guiding 5,000 students to '
    'successful enrolment by 2030.'
)
heading('2.3 Legal & Corporate Structure', level=2)
bullet('Register as a private limited liability company (Ltd) with the Corporate Affairs Commission (CAC), Nigeria.')
bullet('Structure: Founder/Managing Director, Admissions Lead, Visa & Compliance Lead, Operations Manager, and Marketing Manager.')
bullet('Registered office: 14 Ikoyi Road, Lekki Phase 1, Lagos, Nigeria.')

# ==================== 3. MARKET ANALYSIS ====================
heading('3. Market Analysis', level=1)
heading('3.1 Industry Overview', level=2)
para(
    'Nigerian students represent one of Africa\u2019s largest outbound study populations. Demand for European '
    'destinations is driven by affordable tuition in Germany and Poland, strong post-study work rights in Ireland, '
    'and English-taught programmes across the EU. The UK and Canada have tightened visa and cost conditions, '
    'pushing more students toward EU alternatives and making this a strategic window for EU-focused consultants.'
)
heading('3.2 Target Market', level=2)
bullet('Undergraduate aspirants (18–24 months planning horizon).')
bullet('Graduates seeking Master\u2019s and professional programmes (8–12 months horizon).')
bullet('Students needing language/foundation pathways before a degree.')
bullet('Doctoral and research candidates requiring supervision-fit and funding support.')
bullet('Geography: Lagos (headquarters), Abuja, Port Harcourt, Ibadan, and diaspora families via online channels.')
heading('3.3 Market Size & Growth', level=2)
para(
    'Nigeria\u2019s outbound student numbers are estimated in the hundreds of thousands annually, with EU-bound '
    'applications growing year over year. Even capturing a small, focused share (0.1–0.3% of the addressable '
    'market) supports the company\u2019s 240-student Year-2 target.'
)
heading('3.4 Competitive Landscape', level=2)
para(
    'Competitors include generalist agencies, free/low-engagement agents, and self-service online platforms. '
    'EU StudyBridge differentiates through: dedicated personal agents for every applicant, transparent pricing '
    'with documented service tiers, EU-country specialisation (6 destinations), integrated visa strategy, and a '
    'digital-first experience (web planner, WhatsApp chat, automated roadmaps).'
)
heading('3.5 SWOT Analysis', level=2)
add_table(
    ['', 'Key Points'],
    [
        ['Strengths', 'Dedicated per-student agents; 6-country EU focus; strong brand and web presence; transparent tiered pricing; 79% interview-request success metric.'],
        ['Weaknesses', 'New entrant without a long track record; dependency on founder relationships with universities; seasonal intake peaks.'],
        ['Opportunities', 'EU visa/cost advantage over UK/Canada; growing English-taught programme supply; scholarship and pathway partnerships; online-first sales reaching diaspora families.'],
        ['Threats', 'Regulatory changes in EU visa policy; currency volatility (₦/€); aggressive competitor pricing; fake-agent reputation risk in the industry.'],
    ],
    widths=[3.5, 12.5],
)
doc.add_paragraph()

# ==================== 4. SERVICES & REVENUE MODEL ====================
heading('4. Services & Revenue Model', level=1)
heading('4.1 Service Lines', level=2)
para(
    'Four study routes mirror the website offering. Each route is delivered by a dedicated agent and ends with '
    'arrival support, so students experience a single accountable point of contact.'
)
add_table(
    ['Route', 'Services Included', 'Avg. Fee (₦)'],
    [
        ['Undergraduate', 'University shortlisting, applications, personal statement, document pack, offer follow-up', '650,000'],
        ['Master\u2019s / Postgraduate', 'SOP/CV review, scholarship matching, applications, interview prep', '600,000'],
        ['Language + Foundation', 'Pathway matching, intensive English support, fast-track enrolment', '500,000'],
        ['Doctoral / Research', 'Supervisor outreach, research proposal review, funding feasibility', '750,000'],
    ],
    widths=[4.0, 9.0, 3.0],
)
doc.add_paragraph()
heading('4.2 Revenue Streams', level=2)
bullet('Per-student service fees (primary, ~85% of revenue).')
bullet('Visa & pre-departure add-on packages (insurance referral, accommodation booking).')
bullet('Commission from partner universities (standard international education practice).')
bullet('Premium tier: fully managed end-to-end service with priority responsiveness and concierge arrival support.')
heading('4.3 Pricing Strategy', level=2)
para(
    'Fees are tiered by route and complexity, benchmarked against generalist agencies while emphasising '
    'transparency and documentation. Installment plans (e.g., 50% on engagement, 30% on application submission, '
    '20% on offer/visa stage) improve affordability and cash-flow predictability.'
)

# ==================== 5. MARKETING & SALES ====================
heading('5. Marketing & Sales Strategy', level=1)
heading('5.1 Positioning', level=2)
para(
    '"Study in Europe without the confusion." EU StudyBridge is the specialist partner for Nigerian students '
    'who want structured, personal guidance to EU universities — differentiated from generalist agents by '
    'country focus and a dedicated agent for every applicant.'
)
heading('5.2 Channels', level=2)
bullet('Digital: SEO on study-in-Europe keywords, Instagram (@eustudybridgenigeria), Facebook, LinkedIn, and Google Ads.')
bullet('Referrals: 20% referral credit for past students and families — the strongest acquisition channel.')
bullet('Schools & partners: career fairs, university partnerships, and alumni ambassador programme.')
bullet('WhatsApp funnel: the website chat and WhatsApp lines capture leads into a structured follow-up cadence.')
heading('5.3 Sales Funnel Targets', level=2)
add_table(
    ['Stage', 'Metric'],
    [
        ['Leads / month (Year 1)', '300'],
        ['Consultation calls', '90'],
        ['Proposals sent', '45'],
        ['Enrolled students', '12–15 / month by Q4'],
        ['Conversion (lead → enrolled)', '4–5%'],
    ],
    widths=[8.0, 8.0],
)
doc.add_paragraph()

# ==================== 6. OPERATIONS PLAN ====================
heading('6. Operations Plan', level=1)
heading('6.1 Student Journey', level=2)
para('A structured 9–12 month pipeline mirrors the website timeline, managed in a CRM with milestone tracking:')
add_table(
    ['Phase', 'Months', 'Key Activities'],
    [
        ['Research & shortlisting', '1–3', 'Course/country selection, university shortlist, budget and scholarship scan'],
        ['Application & documents', '3–6', 'Transcripts, SOP, recommendation letters, English proficiency, financial evidence'],
        ['Offer & scholarship follow-up', '6–9', 'Track applications, respond to admissions, secure offers, finalise funding'],
        ['Visa, accommodation & travel', '9–12', 'Visa submission, accommodation booking, insurance, pre-departure coaching'],
    ],
    widths=[4.5, 2.0, 9.5],
)
doc.add_paragraph()
heading('6.2 Technology Stack', level=2)
bullet('Website with interactive study planner (already built) — lead capture and self-service roadmap.')
bullet('CRM for pipeline, document checklists, and milestone reminders.')
bullet('WhatsApp Business API for communication and support.')
bullet('Document vault for secure uploads of transcripts, passports, and financial evidence.')
heading('6.3 Capacity & Scalability', level=2)
para(
    'Each dedicated agent manages 25–30 active students. With a team of 4 consultants in Year 1 and 8 by the '
    'end of Year 2, the model supports 240+ enrolled students annually without sacrificing personalised service.'
)

# ==================== 7. MANAGEMENT & ORGANIZATION ====================
heading('7. Management & Organization', level=1)
add_table(
    ['Role', 'Responsibility'],
    [
        ['Founder / Managing Director', 'Strategy, partnerships, university relations, funding'],
        ['Lead Admissions Consultant', 'Undergraduate & Master\u2019s admissions, quality control'],
        ['Visa & Travel Specialist', 'Visa strategy, document compliance, pre-departure'],
        ['Operations Manager', 'CRM, scheduling, client experience, team coordination'],
        ['Marketing Manager', 'Lead generation, brand, content, partnerships'],
    ],
    widths=[5.5, 10.5],
)
doc.add_paragraph()
para(
    'The website already names two client-facing leads: Sarah Okafor (Lead Admissions Consultant) and Daniel Eze '
    '(Visa & Travel Specialist), establishing a credible team image from day one.'
)

# ==================== 8. FINANCIAL PLAN ====================
heading('8. Financial Plan', level=1)
heading('8.1 Pricing Summary', level=2)
para('See Appendix A for the full service fee schedule with add-ons.')
heading('8.2 Year-1 Start-Up Costs', level=2)
add_table(
    ['Item', 'Cost (₦)'],
    [
        ['Company registration & legal', '1,000,000'],
        ['Office fit-out & rent deposit (Lekki)', '2,500,000'],
        ['Equipment & IT (laptops, phones, CRM setup)', '2,500,000'],
        ['Website & digital assets', '1,000,000'],
        ['Initial marketing (launch campaign)', '3,000,000'],
        ['Working capital buffer', '5,000,000'],
        ['Miscellaneous / contingency', '1,000,000'],
        ['TOTAL', '16,000,000'],
    ],
    widths=[11.0, 5.0],
)
doc.add_paragraph()
heading('8.3 Monthly Operating Costs (Year 1, steady state)', level=2)
add_table(
    ['Item', 'Monthly Cost (₦)'],
    [
        ['Salaries (5 staff incl. founder draw)', '2,400,000'],
        ['Office rent & utilities', '500,000'],
        ['Marketing (digital + events)', '800,000'],
        ['Software & CRM', '150,000'],
        ['Travel & meetings', '200,000'],
        ['Miscellaneous', '200,000'],
        ['TOTAL', '4,250,000'],
    ],
    widths=[11.0, 5.0],
)
doc.add_paragraph()
heading('8.4 Revenue Projection', level=2)
para(
    'Year 1: ~96 enrolled students \u2192 ≈ ₦64.0m revenue, reaching breakeven in month 7. '
    'Year 2: ~240 enrolled students \u2192 ≈ ₦186.4m revenue, with a 49% net margin after scale-up costs. '
    'See Appendix B for the full 24-month model.'
)
heading('8.5 Key Assumptions', level=2)
bullet('Average fee per student (blended across routes): ₦667,000.')
bullet('Student acquisition cost stabilises at ~₦85,000 per enrolled student by Year 2.')
bullet('Payment is 50% on engagement; balance collected within 6 months, supporting working capital.')
bullet('25% of fees are collected via university commission offsets (margin-neutral).')

# ==================== 9. FUNDING REQUIREMENTS ====================
heading('9. Funding Requirements', level=1)
para(
    'EU StudyBridge seeks ₦18 million in start-up funding: ₦16 million for the capital expenditure and '
    'pre-opening costs in Section 8.2, plus a ₦2 million contingency. Funds will be deployed as follows:'
)
bullet('60% — operations and working capital (salaries, rent, tools) to reach breakeven.')
bullet('25% — brand, marketing, and lead acquisition.')
bullet('15% — registration, equipment, and office set-up.')
para(
    'Investor options: equity (10–15% for ₦18m), a convertible note, or an Islamic-compliant/simple loan with '
    'revenue-share terms. Repayment capacity is supported by positive cash flow from month 7 onward.'
)

# ==================== 10. RISKS & MITIGATION ====================
heading('10. Risks & Mitigation', level=1)
add_table(
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['EU visa policy tightening', 'High', 'Multi-country portfolio; early visa preparation; compliance expertise'],
        ['₦/€ currency volatility', 'Medium', 'Price fees in ₦ with periodic review; forward-planning for clients'],
        ['Seasonal demand spikes', 'Medium', 'Staggered intakes (Sept/Jan/May); scalable agent capacity model'],
        ['Competitor price pressure', 'Medium', 'Differentiate on dedicated agents, transparency, and success metrics'],
        ['Reputation risk (industry-wide fake agents)', 'High', 'CAC registration, published office, verified reviews, strong documentation'],
    ],
    widths=[4.5, 2.0, 9.5],
)
doc.add_paragraph()

# ==================== 11. IMPLEMENTATION ROADMAP ====================
heading('11. Implementation Roadmap', level=1)
add_table(
    ['Period', 'Milestones'],
    [
        ['Months 0–3', 'Register company, finalise office, hire core team, launch marketing campaign, sign first 5 university partnerships'],
        ['Months 4–6', 'First 30 enrolled students; refine CRM and document checklists; establish referral programme'],
        ['Months 7–12', 'Reach 96 enrolled students; breakeven; expand to Abuja and Port Harcourt via online channels'],
        ['Year 2', '240 enrolled students; 8 consultants; 15+ university partnerships; premium tier launched; ₦186m revenue'],
    ],
    widths=[3.5, 12.5],
)
doc.add_paragraph()

# ==================== APPENDIX A ====================
doc.add_page_break()
heading('Appendix A: Service Fee Schedule', level=1)
add_table(
    ['Service', 'Fee (₦)', 'Notes'],
    [
        ['Undergraduate route', '650,000', 'Foundation year, direct, or pathway admissions'],
        ['Master\u2019s / Postgraduate route', '600,000', 'Includes SOP/CV and scholarship matching'],
        ['Language + Foundation route', '500,000', 'Includes intensive English support'],
        ['Doctoral / Research route', '750,000', 'Supervisor outreach + proposal support'],
        ['Visa & pre-departure add-on', '150,000', 'Insurance, accommodation, interview prep'],
        ['Premium managed package', '900,000', 'Priority + concierge arrival support'],
    ],
    widths=[6.5, 3.5, 6.0],
)
doc.add_paragraph()

# ==================== APPENDIX B ====================
heading('Appendix B: 24-Month Financial Projection', level=1)
add_table(
    ['Metric', 'Year 1', 'Year 2'],
    [
        ['Enrolled students', '96', '240'],
        ['Average fee (₦)', '667,000', '690,000'],
        ['Gross revenue (₦)', '64,000,000', '186,400,000'],
        ['Direct service costs', '19,200,000', '46,600,000'],
        ['Operating expenses', '36,000,000', '64,000,000'],
        ['Net profit (₦)', '8,800,000', '75,800,000'],
        ['Net margin', '14%', '41%'],
    ],
    widths=[6.0, 5.0, 5.0],
)
doc.add_paragraph()
para(
    'Note: Year-2 margin rises as marketing cost per student falls and agent capacity doubles with only two '
    'additional hires. Figures are planning estimates and should be revisited quarterly against live data.',
    size=9, italic=True, color=MUTED,
)

doc.save(r'C:\Users\PATRICK ISOLOKWU\Documents\soft\EU_StudyBridge_Business_Plan.docx')
print('Saved: EU_StudyBridge_Business_Plan.docx')
